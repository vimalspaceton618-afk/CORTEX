#!/usr/bin/env python3
"""
Extract Q&A from docx using python-docx to preserve structure.
Looks for paragraphs that start with "Q<number>:" or numbered questions.
"""

import sys
from pathlib import Path
import re

try:
    import docx
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(1)


def extract_questions_from_docx(docx_path: str, output_txt: str):
    """Extract questions from docx, assuming Q<number>: format."""

    print(f"Opening: {docx_path}")
    document = docx.Document(docx_path)

    questions = []
    current_question = None

    # Patterns for question detection
    question_patterns = [
        re.compile(r'^Q\s*(\d+)\s*[:.]\s*(.+)$', re.IGNORECASE),  # Q1: What...
        re.compile(r'^(\d+)\.\s+(.+)$'),  # 1. What...
        re.compile(r'^Question\s+\d+[:.]\s*(.+)$', re.IGNORECASE),  # Question 1: What...
    ]

    answer_pattern = re.compile(r'^A\s*(\d+)\s*[:.]\s*', re.IGNORECASE)  # A1: ...

    print(f"Document has {len(document.paragraphs)} paragraphs")

    for para in document.paragraphs:
        text = para.text.strip()

        if not text:
            continue

        # Check if this is an answer line
        if answer_pattern.match(text):
            # Save pending question before resetting
            if current_question:
                questions.append(current_question)
                current_question = None
            continue

        # Check for question patterns
        for pattern in question_patterns:
            match = pattern.match(text)
            if match:
                # Extract the actual question text
                if len(match.groups()) > 1:
                    question_text = match.group(2).strip()
                else:
                    question_text = text
                current_question = question_text
                break

        # If we're building a multi-line question, continue it
        if current_question and not any(pattern.match(text) for pattern in question_patterns) and not answer_pattern.match(text):
            # Append this line to current question (continuation)
            current_question += ' ' + text

    # Save last question if pending
    if current_question:
        questions.append(current_question)

    # Deduplicate
    unique_questions = list(dict.fromkeys(questions))

    print(f"Found {len(questions)} questions")
    print(f"Unique: {len(unique_questions)}")

    if unique_questions:
        print("\nFirst 5 questions:")
        for i, q in enumerate(unique_questions[:5], 1):
            print(f"  {i}. {q[:100]}...")

    # Write to file
    with open(output_txt, 'w', encoding='utf-8') as f:
        for q in unique_questions:
            f.write(q + '\n')

    print(f"\nSaved {len(unique_questions)} questions to: {output_txt}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python extract_qa_docx.py <input.docx> <output.txt>")
        sys.exit(1)

    extract_questions_from_docx(sys.argv[1], sys.argv[2])
