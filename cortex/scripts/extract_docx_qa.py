#!/usr/bin/env python3
"""
Extract Q&A from DOCX file (AI_Training_QA_Dataset_100K.docx)
"""

import sys
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)


def extract_qa_from_docx(docx_path: str, output_path: str = None):
    """
    Extract Q&A pairs from DOCX.
    Assumes format: Question on one line, Answer on next line (or separated by pattern).
    """
    doc = Document(docx_path)
    qa_pairs = []

    print(f"[EXTRACT] Reading {docx_path}...")
    print(f"[EXTRACT] Total paragraphs: {len(doc.paragraphs)}")

    # Heuristic: Q&A patterns
    current_question = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this looks like a question
        if text.endswith('?') or text.startswith('What') or text.startswith('How') or text.startswith('Explain') or text.startswith('Why') or text.startswith('Can'):
            # Save previous if exists
            if current_question and current_question.get('question'):
                qa_pairs.append(current_question)
            current_question = {'question': text, 'answer': ''}
        elif current_question and not current_question.get('answer'):
            # This is the answer
            current_question['answer'] = text
        elif current_question and current_question.get('answer'):
            # Multi-line answer - append
            current_question['answer'] += '\n' + text

    # Don't forget last
    if current_question and current_question.get('question'):
        qa_pairs.append(current_question)

    print(f"[EXTRACT] Found {len(qa_pairs)} Q&A pairs")

    # Save to JSONL
    if output_path:
        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, 'w', encoding='utf-8') as f:
            for qa in qa_pairs:
                f.write(json.dumps(qa, ensure_ascii=False) + '\n')
        print(f"[EXTRACT] Saved to {output_path}")

    return qa_pairs


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('docx_path', help='Path to DOCX file')
    parser.add_argument('--output', default='training_data_from_docx.jsonl', help='Output JSONL path')
    args = parser.parse_args()

    import json
    pairs = extract_qa_from_docx(args.docx_path, args.output)

    # Show sample
    if pairs:
        print("\nSample Q&A:")
        print(f"Q: {pairs[0]['question'][:100]}")
        print(f"A: {pairs[0]['answer'][:100]}...")
